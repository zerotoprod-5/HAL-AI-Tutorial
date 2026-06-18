"""
Markdown -> .docx for the planning docs, so they upload cleanly into Google Docs
(Drive -> Open with Google Docs) with headings, tables, bold/italic preserved.
Handles the subset these docs use: # / ## headings, GFM tables, '-' bullet lists,
**bold** / *italic* / `code`, paragraphs, and '---' rules. Not a full Markdown parser.

    python3 build/md_to_docx.py SESSION_PLAN.md "Predictive AI - Session Plan.docx"
"""
import sys, re
from docx import Document
from docx.shared import Pt

def add_runs(paragraph, text):
    """Add text to a paragraph, honouring **bold**, *italic*, `code`."""
    for tok in re.split(r'(\*\*.+?\*\*|\*.+?\*|`.+?`)', text):
        if not tok:
            continue
        if tok.startswith('**') and tok.endswith('**'):
            paragraph.add_run(tok[2:-2]).bold = True
        elif tok.startswith('*') and tok.endswith('*'):
            paragraph.add_run(tok[1:-1]).italic = True
        elif tok.startswith('`') and tok.endswith('`'):
            r = paragraph.add_run(tok[1:-1]); r.font.name = 'Consolas'; r.font.size = Pt(9.5)
        else:
            paragraph.add_run(tok)

def is_table_sep(line):
    s = line.strip()
    return bool(s) and set(s) <= set('|:- ') and '-' in s

def split_row(line):
    s = line.strip()
    if s.startswith('|'): s = s[1:]
    if s.endswith('|'): s = s[:-1]
    return [c.strip() for c in s.split('|')]

def convert(md_path, docx_path):
    lines = open(md_path, encoding='utf-8').read().split('\n')
    doc = Document()
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(10.5)

    i, first_heading = 0, True
    while i < len(lines):
        line = lines[i]
        if not line.strip():
            i += 1; continue

        if line.startswith('## '):
            h = doc.add_heading(level=2); add_runs(h, line[3:].strip()); i += 1; continue
        if line.startswith('# '):
            txt = line[2:].strip()
            if first_heading:
                t = doc.add_heading(level=0); add_runs(t, txt); first_heading = False
            else:
                h = doc.add_heading(level=1); add_runs(h, txt)
            i += 1; continue

        if line.lstrip().startswith('|'):
            block = []
            while i < len(lines) and lines[i].lstrip().startswith('|'):
                block.append(lines[i]); i += 1
            rows = [r for r in block if not is_table_sep(r)]
            if not rows:
                continue
            cells0 = split_row(rows[0]); ncol = len(cells0)
            table = doc.add_table(rows=0, cols=ncol)
            try: table.style = 'Light Grid Accent 1'
            except Exception: table.style = 'Table Grid'
            for ridx, r in enumerate(rows):
                cells = split_row(r)
                cells += [''] * (ncol - len(cells))
                row = table.add_row().cells
                for cidx in range(ncol):
                    p = row[cidx].paragraphs[0]
                    add_runs(p, cells[cidx])
                    if ridx == 0:
                        for run in p.runs: run.bold = True
            doc.add_paragraph()
            continue

        if line.strip() == '---':
            i += 1; continue

        if re.match(r'^\s*-\s+', line):
            while i < len(lines) and re.match(r'^\s*-\s+', lines[i]):
                item = re.sub(r'^\s*-\s+', '', lines[i]); i += 1
                # gather wrapped continuation lines
                while i < len(lines) and lines[i].strip() and not re.match(r'^\s*-\s+', lines[i]) \
                        and not lines[i].startswith('#') and not lines[i].lstrip().startswith('|'):
                    item += ' ' + lines[i].strip(); i += 1
                p = doc.add_paragraph(style='List Bullet'); add_runs(p, item)
            continue

        # paragraph (gather wrapped lines)
        para = line.strip(); i += 1
        while i < len(lines) and lines[i].strip() and not lines[i].startswith('#') \
                and not lines[i].lstrip().startswith('|') and not re.match(r'^\s*-\s+', lines[i]) \
                and lines[i].strip() != '---':
            para += ' ' + lines[i].strip(); i += 1
        p = doc.add_paragraph(); add_runs(p, para)

    doc.save(docx_path)
    print('wrote', docx_path)

if __name__ == '__main__':
    convert(sys.argv[1], sys.argv[2])
